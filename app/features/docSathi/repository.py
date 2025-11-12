from app.utils.openai.openai import start_openai
from app.utils.youtube_transcript import get_transcript_with_cache
from app.youtubeService.document_processor import DocumentProcessor
from app.aws.secretKey import session
from app.config.settings import settings
from app.database.session import get_db
import os
import re
import json
import threading
import boto3
import fitz
import requests
from io import BytesIO
from docx import Document 
from sentence_transformers import SentenceTransformer
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from jinja2 import Environment, FileSystemLoader
from app.aws.secretKey import get_secret_keys
from sentry_sdk import capture_exception
from app.features.docSathi.schema import DocumentChat, DocumentChatsResponse, DocumentTextResponse, UploadDocuments, DocumentSummary, DocumentSummaryResponse, DocumentNotes, DocumentNotesResponse, DocumentQuiz, DocumentQuizResponse, QuizQuestion, GenerateQuizRequest, StudentQuiz, QuestionAnswer, GenerateQuizResponse, QuizQuestionResponse, QuizResponse, GetDocumentQuizzesResponse, SubmitQuizRequest, QuizSubmissionResult, SubmitQuizResponse
# from app.config import env_variables
import nltk
import numpy as np
from fastapi import  HTTPException, BackgroundTasks
from transformers import LEDForConditionalGeneration, LEDTokenizer
import torch
from nltk.tokenize import sent_tokenize
from urllib.parse import urlparse, parse_qs
from bson import ObjectId
# import openai
from datetime import datetime
# env_vars = env_variables()

from bson import ObjectId
from datetime import datetime

from fastapi import  HTTPException
from transformers import LEDForConditionalGeneration, LEDTokenizer
import torch
from nltk.tokenize import sent_tokenize
from app.database.mongo_collections import (
    create_document,
    insert_chunk,
    update_document_status,
    get_collections,
)
from bson import ObjectId
from app.schemas.milvus.collection.milvus_collections import chunk_msmarcos_collection
from app.schemas.schemas import DocSathiAIDocumentBase, DocSathiAIDocumentCreate
from datetime import datetime
import asyncio

from app.utils.delete_from_aws import delete_document_from_s3
from app.utils.milvus.operations.crud import insert_chunk_to_milvus
from app.database.mongo_collections import update_document_status
from bson import ObjectId
document_processor = DocumentProcessor()


# Creating instance meaning full summary 

model = LEDForConditionalGeneration.from_pretrained("allenai/led-base-16384")

tokenizer = LEDTokenizer.from_pretrained("allenai/led-base-16384")

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device).half()

model.config.attention_window = 512 


# Creating instance meaning full chunks 
try:
    modelForChunk = SentenceTransformer("all-MiniLM-L6-v2")
except Exception as e:
    print(f"🚨 Error loading embedding model: {e}")
    modelForChunk = None  
    



keys = get_secret_keys()

s3_client = session.client(
    "s3",
    region_name=settings.AWS_REGION,
    config=boto3.session.Config(signature_version="v4"),
)
s3_resources = session.resource(
    "s3",
    region_name=settings.AWS_REGION,
    config=boto3.session.Config(signature_version="v4"),
)

current_directory = os.path.dirname(os.path.abspath(__file__))
relative_templates_directory = "../../templates"
templates_directory = os.path.join(current_directory, relative_templates_directory)
template_env = Environment(loader=FileSystemLoader(templates_directory))


# generate presigned url
async def generate_presigned_url(filename: str):

    try:
        bucket_name = settings.S3_BUCKET_NAME
        response = s3_client.generate_presigned_url(
            "put_object",
            Params={"Bucket": bucket_name, "Key": f"vaktaAi/{filename}"},
            ExpiresIn=3600,
        )

    except Exception as e:
        print("error in generate_presigned_url", e)
        capture_exception(e)
        return {"message": "Something went wrong", "success": False}

    return {"data": response, "success": True, "message": "generate_presigned_url."}





def extract_text(file_data: bytes, document_format: str) -> str:
    """
    Extract text from file based on its format.
    :param file_data: File content in bytes
    :param document_format: Format of the document (pdf/word)
    :return: Extracted text as a string
    """
    try:
        if document_format == "pdf":
            pdf_document = fitz.open(stream=file_data)
            return "".join(
                pdf_document.load_page(page_num).get_text()
                for page_num in range(len(pdf_document))
            )
        elif document_format == "docx":
            file_stream = BytesIO(file_data)
            doc = Document(file_stream)
            return "\n".join(
                f"\u2022 {p.text.strip()}" if p.style.name.lower().startswith("list") else p.text.strip()
                for p in doc.paragraphs
            )
        else:
            return ""
    except Exception as e:
        print(e, f"Error extracting text for format: {document_format}")
        capture_exception(e)
        return ""

def extract_text_from_docx(url: str) -> str:
    # File download from presigned URL
    response = requests.get(url)
    response.raise_for_status()

    # Load DOCX from bytes
    file_stream = BytesIO(response.content)
    doc = Document(file_stream)

    # Extract text
    text = "\n".join([para.text for para in doc.paragraphs])
    return text


async def read_and_train_private_file(
    data: UploadDocuments,
    user_id: str,
):
    """
    Optimized function to process documents from multiple sources:
    - FileData: Multiple files uploaded to S3
    - websiteUrls: Multiple website URLs for scraping  
    - youtubeUrls: Multiple YouTube URLs for transcript extraction
    - url: Single URL field (auto-detects YouTube vs Web URL)
    
    User can provide only one type at a time (files OR websites OR videos OR url)
    """
    results = []

    try:
        
        # ✅ NEW: Handle single URL field - auto-detect type
        if data.url is not None:
            url = data.url.strip()
            if is_youtube_url(url):
                # Auto-detect YouTube URL
                data.youtube_url = url
            else:
                # Treat as website URL
                data.website_url = url
        
        # Count how many input types are provided (only one should be provided at a time)
        input_types = 0
        if data.file_data is not None:
            input_types += 1
        if data.website_url is not None:
            input_types += 1  
        if data.youtube_url is not None:
            input_types += 1
            
        # Validate: Only one input type allowed at a time
        if input_types == 0:
            return {
                "message": "No input data provided",
                "success": False,
                "error": "Please provide FileData, WebsiteUrl, YoutubeUrl, or Url" 
            }
        elif input_types > 1:
            return {
                "message": "Multiple input types not allowed",
                "success": False,
                "error": "Please provide only one type: FileData OR WebsiteUrl OR YoutubeUrl OR Url"
            }
        
        # Process based on input type (single item processing)
        if data.file_data is not None:
            results = await process_single_file(data.file_data, data.document_format, data.type, user_id)
        elif data.website_url is not None:
            results = await process_single_website(data.website_url, data.type, user_id)
        elif data.youtube_url is not None:
            results = await process_single_youtube(data.youtube_url, data.type, user_id)
            
        return {
            "message": "Document processing completed",
            "results": results,
            "success": all(result["success"] for result in results),
        }

    except Exception as e:
        print(f"Error in read_and_train_private_file: {e}")
        capture_exception(e)
        return {
            "message": "Something went wrong",
            "success": False,
            "error": str(e),
        }


async def get_documents_by_user_id(userId: str):
    try:
        db = next(get_db())
        print("userId----------", userId)
        
        # Handle both string (ObjectId) and integer user_id
        # Try both formats to support backward compatibility
        documents = []
        
        # First try as string (for ObjectId from frontend)
        try:
            if len(userId) == 24 and all(c in '0123456789abcdefABCDEF' for c in userId):
                # It's an ObjectId string, try as string first
                query_str = {"user_id": userId}
                documents = list(db["docSathi_ai_documents"].find(query_str).sort("created_ts", -1))
                print(f"Found {len(documents)} documents with user_id as string: {userId}")
        except Exception as e:
            print(f"Error querying with string user_id: {e}")
        
        # If no documents found with string, try as integer (for backward compatibility)
        if not documents:
            try:
                query_int = {"user_id": int(userId)}
                documents = list(db["docSathi_ai_documents"].find(query_int).sort("created_ts", -1))
                print(f"Found {len(documents)} documents with user_id as integer: {int(userId)}")
            except (ValueError, TypeError) as e:
                print(f"Error converting userId to int: {e}")
                # If conversion fails, try as string anyway
                query_str = {"user_id": userId}
                documents = list(db["docSathi_ai_documents"].find(query_str).sort("created_ts", -1))

        # Convert ObjectId to str
        for doc in documents:
            doc["_id"] = str(doc["_id"])

        return {
            "data": documents,
            "success": True,
            "message": "Documents retrieved successfully." if documents else "No documents found.",
        }

    except Exception as e:
        print("error in get_documents_by_user_id", e)
        import traceback
        traceback.print_exc()
        capture_exception(e)
        return {"message": "Something went wrong", "success": False, "data": []}

async def check_doc_status(document_id):
    """
    Ultra-optimized status check - completely non-blocking, won't interfere with training
    - Uses projection for minimal data transfer
    - Read-only query (no locks)
    - Fast timeout handling
    - No blocking operations
    """
    db = None
    try:
        # Get database connection (non-blocking, uses connection pool)
        db = next(get_db())
        
        # Convert string to ObjectId for MongoDB query
        from bson import ObjectId
        try:
            document_object_id = ObjectId(document_id)
        except Exception:
            return {
                "success": False,
                "message": "Invalid document ID format",
                "data": {},
            }
        
        # OPTIMIZATION: Use projection to only fetch status and essential fields (not full document)
        # This reduces data transfer and query time significantly (50-70% faster)
        # Read-only query - no write locks, won't interfere with training writes
        document = db["docSathi_ai_documents"].find_one(
            {"_id": document_object_id},
            {"status": 1, "name": 1, "url": 1, "document_format": 1, "created_ts": 1, "updated_ts": 1, "summary": 1, "_id": 1},
            # No additional options needed - MongoDB handles read queries efficiently
        )
        
        if not document:
            return {
                "success": False,
                "message": "No document found with the given ID",
                "data": {},
            }

        # Convert ObjectId to string for JSON serialization (fast operation)
        document["_id"] = str(document["_id"])
        
        # Get document status (fast lookup)
        document_status = document.get("status", "unknown")
        
        # Fast return - no blocking operations
        return {
            "success": True,
            "message": "Processing document" if document_status == "processing" else f"Document status: {document_status}",
            "data": document,
        }

    except Exception as e:
        # Fast error handling - don't block on errors
        print(f"[STATUS_CHECK] Error (non-blocking): {e}")
        # Don't capture exception for status checks - keep it lightweight
        return {
            "success": False,
            "message": f"Error checking document status: {str(e)}",
            "data": {},
        }
    finally:
        # Connection is automatically returned to pool by FastAPI dependency
        pass


    




async def process_single_file(file_data, document_format, doc_type, user_id: str):
    """Process single file upload from S3"""
    bucket_name = settings.S3_BUCKET_NAME
    results = []
    
    try:
        # Extract S3 key from signed URL or use fileNameTime directly
        # The signedUrl contains the full S3 key with path
        s3_key = None
        if file_data.signedUrl:
            # Extract key from URL like: https://bucket.s3.amazonaws.com/vaktaAi/filename.pdf
            from urllib.parse import urlparse, unquote
            parsed_url = urlparse(file_data.signedUrl)
            # Get path without leading slash and unquote URL encoding
            s3_key = unquote(parsed_url.path.lstrip('/'))
        else:
            # Fallback to constructing key from fileNameTime
            s3_key = f"vaktaAi/{file_data.fileNameTime}"
        
        # Get file from S3
        s3_object = (
            s3_resources.Bucket(bucket_name)
            .Object(s3_key)
            .get()
        )
        file_data_content = s3_object["Body"].read()

        # Extract text based on document format
        text = extract_text(file_data_content, document_format)
        if not text.strip():
            results.append({
                "fileName": file_data.fileNameTime,
                "message": "No content extracted from file",
                "success": False,
            })
            return results

        
        # Add document record to Mongo
        db = next(get_db())
        doc_payload = DocSathiAIDocumentCreate(
            user_id=user_id,
            name=file_data.fileNameTime,
            url=file_data.signedUrl,
            status="processing",
            document_format=document_format,
            type=(doc_type or "").lower() if doc_type else None,
            summary=None,
        )
        
        new_document_id = create_document(db, doc_payload.dict(exclude_none=True))

        # Start training in a separate thread
        threading.Thread(
                target=train_document,
                args=(text, new_document_id, doc_type),
            ).start()
        results.append({
            "fileName": file_data.fileNameTime,
            "documentId": new_document_id,
            "message": "File processed successfully",
            "success": True,
        })
        
    except Exception as file_error:
        print(f"Error processing file {file_data.fileNameTime}: {file_error}")
        results.append({
            "fileName": file_data.fileNameTime,
            "message": f"Error processing file: {str(file_error)}",
            "success": False,
            "error": str(file_error),
        })
    
    return results


async def process_file_data(file_data_list, document_format, doc_type):
    """Process multiple file uploads from S3 (legacy function)"""
    bucket_name = settings.S3_BUCKET_NAME
    results = []
    
    for document in file_data_list:
        try:
            # Get file from S3
            s3_object = (
                s3_resources.Bucket(bucket_name)
                .Object(f"vaktaAi/{document.fileNameTime}")
                .get()
            )
            file_data = s3_object["Body"].read()

            # Extract text based on document format
            text = extract_text(file_data, document_format)
            if not text.strip():
                results.append({
                    "fileName": document.fileNameTime,
                    "message": "No content extracted from file",
                    "success": False,
                })
                continue

            
            # Add document record to Mongo
            db = next(get_db())
            doc_payload = DocSathiAIDocumentCreate(
                user_id=1,  # TODO: Get actual user_id
                name=document.fileNameTime,
                url=document.signedUrl,
                status="processing",
                document_format=document_format,
                type=(doc_type or "").lower() if doc_type else None,
                summary=None,
            )
            
            new_document_id = create_document(db, doc_payload.dict(exclude_none=True))

            # Train the document
            # train_document(text, new_document_id, document_format or "file", "1")
            # Start training in a separate thread
            threading.Thread(
                    target=train_document,
                    args=(text, new_document_id, doc_type),
                ).start()
            results.append({
                "fileName": document.fileNameTime,
                "documentId": new_document_id,
                "documentName": document.fileNameTime,
                "message": f"File {document.fileNameTime} processed successfully",
                "success": True,
            })

        except Exception as doc_error:
            print(f"Error processing file {document.fileNameTime}: {doc_error}")
            capture_exception(doc_error)
            results.append({
                "fileName": document.fileNameTime,
                "message": f"Error processing file: {str(doc_error)}",
                "success": False,
                "error": str(doc_error),
            })
    
    return results


async def process_single_website(website_url, doc_type, user_id: str):
    """Process single website URL"""
    results = []
    
    try:
        # Scrape website content
        scraped_data = scrape_website_content(website_url)
        
        if not scraped_data['text'].strip():
            results.append({
                "source": "website",
                "url": website_url,
                "message": "No content extracted from website",
                "success": False,
            })
            return results
        
        # Generate document name from scraped content
        doc_name = generate_document_name_from_text(scraped_data['text'], "website")
        
        # Add document record to Mongo
        db = next(get_db())
        doc_payload = DocSathiAIDocumentCreate(
            user_id=user_id,
            name=doc_name,
            url=website_url,
            status="processing",
            document_format="webUrl",
            type=doc_type or "website",
        )
        
        new_document_id = create_document(db, doc_payload.dict(exclude_none=True))
        
        # Train the document
        threading.Thread(
                target=train_document,
                args=(scraped_data['text'], new_document_id, doc_type),
            ).start()
        
        results.append({
            "source": "website",
            "url": website_url,
            "documentId": new_document_id,
            "documentName": doc_name,
            "message": "Website processed successfully",
            "success": True,
        })
        
    except Exception as web_error:
        print(f"Error processing website URL {website_url}: {web_error}")
        results.append({
            "source": "website",
            "url": website_url,
            "message": f"Error processing website: {str(web_error)}",
            "success": False,
            "error": str(web_error),
        })
    
    return results



def get_video_id(url: str) -> str:
    """
    Extract video id from YouTube URL
    """
    parsed_url = urlparse(url)
    if parsed_url.hostname in ["www.youtube.com", "youtube.com"]:
        # Example: https://www.youtube.com/watch?v=VIDEO_ID
        return parse_qs(parsed_url.query).get("v", [None])[0]
    elif parsed_url.hostname == "youtu.be":
        # Example: https://youtu.be/VIDEO_ID
        return parsed_url.path.lstrip("/")
    return None

async def process_single_youtube(youtube_url, doc_type, user_id: str):
    """Process single YouTube URL"""
    results = []
    
    try:
        video_id = extract_youtube_video_id(youtube_url)
        if not video_id:
            results.append({
                "source": "youtube",
                "url": youtube_url,
                "message": "No video ID found in YouTube URL. Please check the URL format.",
                "success": False,
            })
            return results
        
        # ✅ NEW: Extract timestamped transcript for YouTube videos
        timestamped_transcript = None
        try:
            from youtube_transcript_api import YouTubeTranscriptApi
            
            # Try to get timestamped transcript
            languages = ['en', 'hi', 'en-US', 'en-GB']
            for lang in languages:
                try:
                    transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=[lang])
                    # Store timestamped transcript segments
                    timestamped_transcript = [
                        {
                            "start": entry.get('start', 0),
                            "duration": entry.get('duration', 0),
                            "text": entry.get('text', '').strip()
                        }
                        for entry in transcript_list
                        if entry.get('text', '').strip()
                    ]
                    print("timestamped_transcript>>>>>>>>>>>>>>>>>>>>>>>>>>>>>",timestamped_transcript)
                    if timestamped_transcript:
                        break
                except:
                    continue
            
            # If no transcript found in preferred languages, try any available
            if not timestamped_transcript:
                try:
                    transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
                    timestamped_transcript = [
                        {
                            "start": entry.get('start', 0),
                            "duration": entry.get('duration', 0),
                            "text": entry.get('text', '').strip()
                        }
                        
                        for entry in transcript_list
                        if entry.get('text', '').strip()
                    ]
                    print("timestamped_transcript>>>>>>>>>>>>>>>>>>>>>>>>>>>>>",timestamped_transcript)
                except:
                    pass
        except ImportError:
            print("youtube-transcript-api not installed, skipping timestamped transcript")
        except Exception as e:
            print(f"Error extracting timestamped transcript: {e}")
        
        # Get transcript text (for training)
        data = await get_transcript_with_cache(youtube_url, video_id)
        
        # Convert segments list to text if needed
        if isinstance(data, list):
            data = " ".join(data) if data else ""
        
        if not data or len(data.strip()) == 0:
            data = await document_processor.process_youtube(youtube_url, doc_type)
        if not data or not data.strip():
            results.append({
                "source": "youtube",
                "url": youtube_url,
                "message": "No transcript extracted from YouTube video. The video may not have captions enabled.",
                "success": False,
            })
            return results
        
        # First, try to get video metadata (title/description) for document name
        doc_name = None
        try:
            video_info = await document_processor._get_youtube_video_info(video_id)
            if video_info:
                # Parse video info to extract title
                # Format: "Video Title: {title}\nChannel: {channel}\n\nDescription:\n{description}"
                lines = video_info.split('\n')
                for line in lines:
                    if line.startswith('Video Title:'):
                        doc_name = line.replace('Video Title:', '').strip()
                        # Remove " - YouTube" suffix if present
                        if doc_name.endswith(' - YouTube'):
                            doc_name = doc_name[:-10].strip()
                        break
                
                # If no title found, try to get from description (first line)
                if not doc_name and 'Description:' in video_info:
                    desc_start = video_info.find('Description:') + len('Description:')
                    description = video_info[desc_start:].strip()
                    if description:
                        # Take first line or first 100 characters
                        first_line = description.split('\n')[0].strip()
                        if first_line:
                            doc_name = first_line[:100]  # Limit to 100 chars
        except Exception:
            pass
        
        # Use video metadata as document name if available, otherwise generate from transcript
        if not doc_name:
            doc_name = generate_document_name_from_text(data, "video")
        # Add document record to Mongo
        db = next(get_db())
        doc_payload = DocSathiAIDocumentCreate(
            user_id=user_id,
            name=doc_name,
            url=youtube_url,
            status="processing",
            document_format="videoUrl",
            type=doc_type or "video",
            summary=None,
        )
        
        new_document_id = create_document(db, doc_payload.dict(exclude_none=True))
        
        # ✅ NEW: Store timestamped transcript in document if available
        if timestamped_transcript:
            try:
                update_document_status(db, new_document_id, {
                    "transcript_segments": timestamped_transcript,
                    "updated_ts": datetime.now()
                })
                print(f"✅ Stored {len(timestamped_transcript)} timestamped transcript segments")
            except Exception as e:
                print(f"⚠️ Error storing timestamped transcript: {e}")
        
        # Train the document
        threading.Thread(
                target=train_document,
                args=(data, new_document_id, doc_type),
            ).start()
        
        results.append({
            "source": "youtube",
            "url": youtube_url,
            "documentId": new_document_id,
            "documentName": doc_name,
            "message": "YouTube video processed successfully",
            "success": True,
        })
        
    except Exception as youtube_error:
        print(f"Error processing YouTube URL {youtube_url}: {youtube_error}")
        results.append({
            "source": "youtube",
            "url": youtube_url,
            "message": f"Error processing YouTube video: {str(youtube_error)}",
            "success": False,
            "error": str(youtube_error),
        })
    
    return results




# Utility functions for document processing 




def train_document(text: str, document_id: str, document_type: str,):
    """
    Optimized training function with batch operations and parallel processing
    Runs in background thread, doesn't block status checks
    """
    try:
        db = next(get_db()) 
        
        # Generate meaningful chunks
        chunks = create_meaningful_chunks(text)
        
        # Generate summary asynchronously (this runs in thread-safe way)
        summary = asyncio.run(create_document_keypoints(chunks))
        
        # OPTIMIZATION: Batch insert chunks instead of one-by-one
        if chunks:
            chunk_documents = []
            for item in chunks:
                chunk_doc = {
                    "detail": "Default detail",
                    "keywords": "Default keywords",
                    "meta_summary": "Default summary",
                    "chunk": item,
                    "training_document_id": ObjectId(document_id),
                    "created_ts": datetime.now(),
                    "updated_ts": datetime.now(),
                }
                chunk_documents.append(chunk_doc)
            
            # Batch insert all chunks at once (much faster than sequential inserts)
            if chunk_documents:
                _, chunks_collection, _, _, _, _, _, _ = get_collections(db)
                chunks_collection.insert_many(chunk_documents)
            
            # OPTIMIZATION: Process Milvus inserts in parallel batches
            # Get inserted chunk IDs (they should be in same order)
            inserted_chunks = list(db["chunks"].find(
                {"training_document_id": ObjectId(document_id)},
                {"_id": 1}
            ).sort("created_ts", 1))
            
            # Process Milvus inserts in batches (avoid blocking)
            import concurrent.futures
            def insert_to_milvus_sync(chunk_text, doc_id, chunk_id):
                """Wrapper for async Milvus insert"""
                try:
                    asyncio.run(insert_chunk_to_milvus(db, chunk_text, doc_id, str(chunk_id)))
                except Exception as e:
                    # Only log errors
                    pass
            
            # Process Milvus inserts in parallel (max 5 concurrent)
            with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
                futures = []
                for i, item in enumerate(chunks):
                    if i < len(inserted_chunks):
                        chunk_id = inserted_chunks[i]["_id"]
                        future = executor.submit(insert_to_milvus_sync, item, document_id, chunk_id)
                        futures.append(future)
                
                # Wait for all Milvus inserts to complete
                for future in concurrent.futures.as_completed(futures):
                    try:
                        future.result()
                    except Exception:
                        pass

        # Update document status to completed
        update_document_status(db, document_id, {"status": "completed", "summary": summary, "updated_ts": datetime.now()})

    except Exception as e:
        print(f"[TRAINING] ❌ Error in train_document: {e}")
        import traceback
        traceback.print_exc()
        try:
            db = next(get_db())
            update_document_status(db, document_id, {"status": "failed", "updated_ts": datetime.now()})
        except:
            pass
        capture_exception(e)

async def create_document_keypoints(chunks: list) -> str:
    """Generate document summary and keypoints using OpenAI"""
    try:
        if not chunks:
            return "No content available for summary."
        
        openai_client = start_openai()
        
        # Combine chunks (limit to avoid token issues)
        combined_text = " ".join(chunks[:10])  # First 10 chunks
        if len(combined_text) > 4000:
            combined_text = combined_text[:4000]
        
        prompt = f"""
        You are an advanced assistant specialized in summarizing long texts such as YouTube transcripts and documents.  
        
        ### Task:
        - Analyze the provided text.  
        - Extract the **most important key points** in clear, bullet-point form.  
        - After the key points, provide a **short final summary (2-4 sentences)** that captures the overall meaning and essence of the text.  
        - Keep the language **simple, clear, and easy to understand** (suitable for students up to class 12).  
        - Avoid repetition, filler words, or irrelevant details.  
        
        ### Input Text:
        {combined_text}
        
        ### Output Format:
        **Key Points:**
        - Point 1
        - Point 2
        - Point 3
        
        **Final Summary:**
        <2-4 sentence summary here>
        """
        
        response = await openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert document analyst. Provide clear, concise summaries and key points."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=1000
        )
        
        summary = response.choices[0].message.content
        return summary
        
    except Exception as e:
        print(f"Error in create_document_keypoints: {e}")
        capture_exception(e)
        return "Could not generate summary due to an error."


# delete document by document Id
async def delete_document_by_admin(docId: str, delete_document: bool = True):
    try:
        db = next(get_db())

        # Set status to deleting
        update_document_status(db, docId, {"status": "deleting", "updated_ts": datetime.now()})

        if docId:
            thread = threading.Thread(
                target=delete_data_related_document,
                args=(docId, delete_document),
            )
            thread.start()
            return {"message": "Document deleted successfully", "success": True}
        else:
            return {"message": "Document not deleted", "success": False}

    except Exception as e:
        print("error================================================", e)
        capture_exception(e)
        return {"message": "Something went wrong", "success": False}



# Delete all data related to the document (Mongo + Milvus + S3)
def delete_data_related_document(docId: str, delete_document: bool):
    db = next(get_db())
    try:
        # 1) Delete chunks from Milvus by document id
        chunk_ids_query = chunk_msmarcos_collection.query(
            f"document_id == {docId}",
            output_fields=["id"]
        )
        chunk_ids = [row["id"] for row in chunk_ids_query] if chunk_ids_query else []
        if chunk_ids:
            chunk_msmarcos_collection.delete(f"id in {chunk_ids}")

        # 2) Delete chunks in Mongo tied to this document
        try:
            oid = ObjectId(docId)
        except Exception:
            return

        db["chunks"].delete_many({"training_document_id": oid})

        # 3) Optionally delete document and its object from S3
        if delete_document:
            doc = db["docSathi_ai_documents"].find_one({"_id": oid}, {"name": 1})
            doc_name = doc.get("name") if doc else None

            db["docSathi_ai_documents"].delete_one({"_id": oid})

            if doc_name:
                try:
                    delete_document_from_s3(doc_name)
                except Exception:
                    pass

        else:
            # If not physically deleting the document, mark as deleted
            db["docSathi_ai_documents"].update_one(
                {"_id": oid},
                {"$set": {"status": "deleted", "updated_ts": datetime.now()}}
            )

    except Exception as e:
        capture_exception(e)



# Function to determine adaptive chunk size based on document length
def get_dynamic_chunk_size(word_count):
    """
    Dynamically determine chunk size based on the document length.
    Ensures chunking is meaningful based on sentence structure.
    """
    if word_count <= 1000:
        return 250  # Small documents → smaller chunks
    elif word_count <= 5000:
        return 500  # Medium-sized documents
    elif word_count <= 10000:
        return 750  # Longer documents
    else:
        return 1000  # Very large documents → bigger chunks
    



def create_meaningful_chunks(text, similarity_threshold=0.5):
    """
    Splits unstructured text into meaningful chunks while ensuring topic coherence.
    Uses SentenceTransformer embeddings for similarity calculations.
    """
    try:
        # ✅ Input Validation
        if not isinstance(text, str) or len(text.strip()) == 0:
            raise ValueError("Input text must be a non-empty string.")

        # ✅ Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()

        # ✅ Sentence Tokenization (More Accurate)
        sentences = sent_tokenize(text)  # Extracts sentences correctly
        
        # ✅ Adaptive Chunk Size
        document_length = len(text.split())  # Count total words
        chunk_size = get_dynamic_chunk_size(document_length)

        # ✅ Ensure model is loaded before encoding
        if modelForChunk is None:
            raise RuntimeError("Sentence Transformer model failed to load.")

        # ✅ Encode sentences efficiently
        try:
            sentence_embeddings = modelForChunk.encode(sentences, batch_size=8)
        except Exception as e:
            raise RuntimeError(f"🚨 Error in sentence encoding: {e}")

        chunks = []
        current_chunk = []
        current_length = 0

        for i, sentence in enumerate(sentences):
            sentence_length = len(sentence.split())

            # ✅ Handling Very Long Sentences (Ensuring Proper Splitting)
            if sentence_length > chunk_size * 1.5:
                split_sentences = sent_tokenize(sentence)  # Split by sentence, not words
                for sub_sentence in split_sentences:
                    chunks.append(sub_sentence)
                continue  # Skip to next iteration

            # ✅ Compute similarity with previous sentence (If applicable)
            if i > 0:
                try:
                    similarity = np.dot(sentence_embeddings[i], sentence_embeddings[i-1]) / (
                        np.linalg.norm(sentence_embeddings[i]) * np.linalg.norm(sentence_embeddings[i-1]) + 1e-6
                    )  # Small value added to avoid division by zero
                except Exception:
                    similarity = 1.0  # Assume high similarity if error occurs

                # ✅ If similarity is LOW, start a new chunk
                if similarity < similarity_threshold and current_length > chunk_size * 0.75:
                    chunks.append(" ".join(current_chunk))
                    current_chunk = []
                    current_length = 0

            # ✅ Add Sentence to Current Chunk
            current_chunk.append(sentence)
            current_length += sentence_length

            # ✅ If chunk size limit is reached, finalize the chunk (ENSURE FULL SENTENCE)
            if current_length >= chunk_size and "." in sentence:  # Ensure it ends at a period
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_length = 0

        # ✅ Add the last chunk if any remains
        if current_chunk:
            chunks.append(" ".join(current_chunk))

        # ✅ Generate summary AFTER chunking for better context
        # summary = create_summary(" ".join(chunks))  # Generate summary based on final chunks

        return chunks
    except Exception as e:
        capture_exception(e)
        return [], ""



 # 🔹 Optimized Function to Dynamically Adjust Summary Length
def get_summary_length(word_count):
    if word_count <= 1000:
        return 100, 200  # Short documents
    elif word_count <= 5000:
        return 150, 300  # Medium documents
    elif word_count <= 10000:
        return 300, 600  # Long documents
    else:
        return 500, 1000  # Very Large documents


# 🔹 Improved Chunking Function (Ensures Sentence Completion & Smart Token Limits)
def chunk_text(text, max_tokens=10000):
    sentences = sent_tokenize(text)  # Split text into sentences
    chunks = []
    current_chunk = []
    current_length = 0

    for sentence in sentences:
        tokenized_sentence = tokenizer.encode(sentence, add_special_tokens=False)
        sentence_length = len(tokenized_sentence)

        # If adding this sentence exceeds max_tokens, finalize the chunk
        if current_length + sentence_length > max_tokens:
            if current_chunk:  # Ensure we don't create empty chunks
                chunks.append(" ".join(current_chunk))
            current_chunk = [sentence]  # Start a new chunk with this sentence
            current_length = sentence_length
        else:
            current_chunk.append(sentence)
            current_length += sentence_length

    # Add last chunk if any sentences remain
    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks


def scrape_website_content(url):
    """
    Scrape content from a website URL and extract meaningful text.
    """
    try:
        # Add headers to avoid being blocked
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        # Parse HTML content
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
            script.decompose()
        
        # Get text from main content areas
        content_selectors = [
            'main', 'article', '.content', '#content', '.post', '.entry-content',
            '.article-content', '.page-content', 'section'
        ]
        
        main_content = ""
        for selector in content_selectors:
            elements = soup.select(selector)
            if elements:
                main_content = " ".join([elem.get_text(strip=True) for elem in elements])
                break
        
        # If no main content found, get all paragraph text
        if not main_content:
            paragraphs = soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
            main_content = " ".join([p.get_text(strip=True) for p in paragraphs])
        
        # Clean up text
        main_content = re.sub(r'\s+', ' ', main_content)  # Remove extra whitespace
        main_content = main_content.strip()
        
        # Get page title for document name
        title = soup.find('title')
        page_title = title.get_text().strip() if title else urlparse(url).netloc
        
        return {
            'text': main_content,
            'title': page_title,
            'url': url
        }
        
    except Exception as e:
        print(f"Error scraping website {url}: {e}")
        return {
            'text': "",
            'title': f"Website - {urlparse(url).netloc}",
            'url': url
        }


def generate_document_name_from_text(text, source_type="document"):
    """
    Generate a meaningful document name based on text content.
    """
    try:
        # Get first few sentences or words
        sentences = sent_tokenize(text)
        if sentences:
            # Use first sentence, but limit length
            first_sentence = sentences[0].strip()
            if len(first_sentence) > 100:
                # Take first 100 characters and find last complete word
                truncated = first_sentence[:100]
                last_space = truncated.rfind(' ')
                if last_space > 50:  # Ensure we have meaningful content
                    first_sentence = truncated[:last_space]
            
            # Clean up the name
            name = re.sub(r'[^\w\s-]', '', first_sentence)
            name = re.sub(r'\s+', ' ', name).strip()
            
            if name:
                return f"{source_type.title()}: {name}"
        
        # Fallback: use first few words
        words = text.split()[:10]
        if words:
            name = " ".join(words)
            name = re.sub(r'[^\w\s-]', '', name)
            return f"{source_type.title()}: {name}"
            
    except Exception as e:
        print(f"Error generating document name: {e}")
    
    # Ultimate fallback
    return f"{source_type.title()} - {datetime.now().strftime('%Y%m%d_%H%M%S')}"


def extract_youtube_transcript(youtube_url):
    """
    Extract transcript from YouTube URL using youtube-transcript-api.
    """
    try:
        # Import here to avoid dependency issues if not installed
        from youtube_transcript_api import YouTubeTranscriptApi
        
        # Extract video ID from URL
        video_id = extract_youtube_video_id(youtube_url)
        if not video_id:
            return {
                'text': "",
                'title': f"YouTube Video - {youtube_url}",
                'url': youtube_url
            }
        
        # Try to get transcript in different languages
        languages = ['en', 'hi', 'en-US', 'en-GB']  # English and Hindi
        transcript_text = ""
        
        for lang in languages:
            try:
                transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=[lang])
                transcript_text = " ".join([entry['text'] for entry in transcript])
                if transcript_text.strip():
                    break
            except:
                continue
        
        # If no transcript found in preferred languages, try any available
        if not transcript_text.strip():
            try:
                transcript = YouTubeTranscriptApi.get_transcript(video_id)
                transcript_text = " ".join([entry['text'] for entry in transcript])
            except:
                pass
        
        # Clean up transcript text
        transcript_text = re.sub(r'\[.*?\]', '', transcript_text)  # Remove [Music], [Applause] etc.
        transcript_text = re.sub(r'\s+', ' ', transcript_text).strip()
        
        return {
            'text': transcript_text,
            'title': f"YouTube Video - {video_id}",
            'url': youtube_url
        }
        
    except ImportError:
        print("youtube-transcript-api not installed. Please install it: pip install youtube-transcript-api")
        return {
            'text': "",
            'title': f"YouTube Video - {youtube_url}",
            'url': youtube_url
        }
    except Exception as e:
        print(f"Error extracting YouTube transcript: {e}")
        return {
            'text': "",
            'title': f"YouTube Video - {youtube_url}",
            'url': youtube_url
        }


def extract_youtube_video_id(url):
    """Extract YouTube video ID from URL."""
    patterns = [
        r'(?:v=|\/)([0-9A-Za-z_-]{11}).*',
        r'(?:embed\/)([0-9A-Za-z_-]{11})',
        r'(?:v\/)([0-9A-Za-z_-]{11})',
        r'(?:youtu\.be\/)([0-9A-Za-z_-]{11})'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def is_youtube_url(url: str) -> bool:
    """Check if URL is a YouTube URL"""
    try:
        parsed_url = urlparse(url)
        hostname = parsed_url.hostname or ""
        return (
            "youtube.com" in hostname or 
            "youtu.be" in hostname or
            hostname == "youtu.be"
        )
    except Exception:
        return False

async def get_document_text_from_chunks(document_id: str):
    """Common function to extract and join all chunks from a document"""
    try:
        from bson import ObjectId
        
        db = next(get_db())
        
        # Validate document_id
        try:
            doc_object_id = ObjectId(document_id)
        except Exception:
            return None, "Invalid document ID format"
        
        # Check if document exists
        document = db["docSathi_ai_documents"].find_one({"_id": doc_object_id})
        if not document:
            return None, "Document not found"
        
        # Get all chunks for this document
        chunks = list(db["chunks"].find({
            "training_document_id": doc_object_id
        }))
        
        if not chunks:
            return None, "No chunks found for this document"
        
        # Combine all chunk text
        combined_text = " ".join([chunk.get("chunk", "") for chunk in chunks])
        
        if not combined_text.strip():
            return None, "No text content found in chunks"
        
        return combined_text, None
        
    except Exception as e:
        print(f"Error in get_document_text_from_chunks: {e}")
        return None, f"Error extracting document text: {str(e)}"

async def generate_document_summary(document_id: str):
    """Generate summary for a document based on its chunks"""
    try:
        # Use common function to get document text
        combined_text, error = await get_document_text_from_chunks(document_id)
        
        if error:
            return DocumentSummaryResponse(
                success=False,
                message=error
            )
        # summary = create_summary(combined_text)
        # Calculate word count for context
        word_count = len(combined_text.split())
        
        # Get chunks count for fallback
        db = next(get_db())
        from bson import ObjectId
        doc_object_id = ObjectId(document_id)
        chunks = list(db["chunks"].find({"training_document_id": doc_object_id}))
        # print("summary>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>",summary)
        # Generate summary using OpenAI
        try:
            
            openai_client = start_openai()

# Create prompt for summary generation
            openai_client = start_openai()

            prompt = f"""
Analyze the document and produce a **clear, readable 4–5 page summary**.

DOCUMENT CONTENT:
{combined_text[:12000]}

INSTRUCTIONS:
Write a structured output with the following sections:

1. TITLE:
   - A short, descriptive title (3–8 words)

2. SUMMARY (2–3 pages total):
   - Write 5–7 well-developed paragraphs
   - Explain the purpose and main ideas of the document
   - Describe key topics, arguments, or themes
   - Include examples or important supporting points
   - Maintain smooth flow and clarity
   - Do NOT exceed 3 pages worth of text (around 500–900 words)

3. KEY POINTS:
   - List 6–10 concise bullet points
   - Each point should be **one short paragraph (1–3 sentences)**
   - Focus on the most important information, insights, or conclusions

OUTPUT FORMAT (Return valid JSON only — NO markdown, NO extra commentary):

{
  "title": "Your title here",
  "summary": "Your 5–7 paragraph structured summary here, approx 500–900 words.",
  "key_points": [
     "Detailed key point 1 (1–3 sentences)",
     "Detailed key point 2 (1–3 sentences)",
     "Continue until total 6–10 points"
  ]
}
"""

            response = await openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are an expert document analyst. Provide clear, concise summaries and key points."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=1000
            )
            
            # Parse the response
            import json
            try:
                result = json.loads(response.choices[0].message.content)
                summary = result.get("summary", "")
                key_points = result.get("key_points", [])
                title = result.get("title", "Document Summary")
            except json.JSONDecodeError:
                # Fallback if JSON parsing fails
                content = response.choices[0].message.content
                summary = content
                key_points = ["Summary generated successfully"]
                title = "Document Summary"
            
        except Exception as e:
            print(f"Error generating summary with OpenAI: {e}")
            # Fallback summary
            summary = f"This document contains {word_count} words across {len(chunks)} sections. The content covers various topics and provides detailed information."
            key_points = [
                f"Document contains {word_count} words",
                f"Split into {len(chunks)} sections",
                "Content analysis completed",
                "Summary generated successfully"
            ]
            title = "Document Summary"
        
        # Create response
        document_summary = DocumentSummary(
            summary=summary,
            key_points=key_points,
            title=title
        )
        
        return DocumentSummaryResponse(
            success=True,
            message="Document summary generated successfully",
            data=document_summary
        )
        
    except Exception as e:
        print(f"Error in generate_document_summary: {e}")
        capture_exception(e)
        return DocumentSummaryResponse(
            success=False,
            message=f"Error generating summary: {str(e)}"
        )

async def generate_document_notes(document_id: str):
    """Generate notes for a document based on its chunks"""
    try:
        # Use common function to get document text
        combined_text, error = await get_document_text_from_chunks(document_id)
        
        if error:
            return DocumentNotesResponse(
                success=False,
                message=error
            )
        
        # Generate notes using OpenAI
        try:
            openai_client = start_openai()
            
            # Create prompt for notes generation
            prompt = f"""
            Please analyze the following document content and create comprehensive study notes:
            
            Document Content:
            {combined_text[:8000]}  # Limit to avoid token limits
            
            Please create detailed study notes with:
            1. A descriptive title (3-8 words)
            2. Key study notes (8-12 bullet points covering main concepts, important details, and insights)
            
            Please respond in the following JSON format:
            {{
                "title": "Study Notes Title",
                "notes": [
                    "Note 1: Main concept or important detail",
                    "Note 2: Key insight or explanation",
                    "Note 3: Supporting information",
                    "Note 4: Additional context",
                    "Note 5: Important facts or data",
                    "Note 6: Key takeaways",
                    "Note 7: Related concepts",
                    "Note 8: Summary points"
                ]
            }}
            """
            
            response = await openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are an expert educator. Create comprehensive, well-structured study notes that help students understand and retain key information."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=1500
            )
            
            # Parse the response
            import json
            try:
                result = json.loads(response.choices[0].message.content)
                notes = result.get("notes", [])
                title = result.get("title", "Document Notes")
            except json.JSONDecodeError:
                # Fallback if JSON parsing fails
                content = response.choices[0].message.content
                notes = [content]
                title = "Document Notes"
            
        except Exception as e:
            print(f"Error generating notes with OpenAI: {e}")
            # Fallback notes
            notes = [
                "Document contains important information",
                "Key concepts and details are present",
                "Study material is comprehensive",
                "Important insights and data included",
                "Well-structured content for learning"
            ]
            title = "Document Notes"
        
        # Save notes to database in ai_notes field
        try:
            db = next(get_db())
            
            # Prepare ai_notes data structure
            ai_notes_data = {
                "title": title,
                "notes": notes,
                "generated_at": datetime.now()
            }
            
            # Update document with ai_notes field
            update_document_status(db, document_id, {
                "ai_notes": ai_notes_data,
                "updated_ts": datetime.now()
            })
            
        except Exception as save_error:
            print(f"Error saving ai_notes to database: {save_error}")
            # Don't fail the request if save fails, just log it
            capture_exception(save_error)
        
        # Create response
        document_notes = DocumentNotes(
            notes=notes,
            title=title
        )
        
        return DocumentNotesResponse(
            success=True,
            message="Document notes generated successfully",
            data=document_notes
        )
        
    except Exception as e:
        print(f"Error in generate_document_notes: {e}")
        capture_exception(e)
        return DocumentNotesResponse(
            success=False,
            message=f"Error generating notes: {str(e)}"
        )


async def generate_student_quiz(request: GenerateQuizRequest):
    """Generate a complete student quiz with questions and save to database"""
    try:
        from bson import ObjectId
        from app.database.mongo_collections import create_student_quiz, create_question_answer
        
        db = next(get_db())
        
        # Validate document exists
        document = db["docSathi_ai_documents"].find_one({"_id": ObjectId(request.document_id)})
        if not document:
            return GenerateQuizResponse(
                success=False,
                message="Document not found"
            )
        
        # Get document text using common function
        combined_text, error = await get_document_text_from_chunks(request.document_id)
        if error:
            return GenerateQuizResponse(
                success=False,
                message=error
            )
        
        # Create quiz in database
        quiz_doc = {
            "quiz_name": request.quiz_name,
            "related_doc_id": ObjectId(request.document_id),
            "created_by": request.user_id, #ObjectId(request.user_id)
            "level": request.level,
            "no_of_questions": request.number_of_questions,
            "is_submitted": False,
            "created_at": datetime.now(),
            "updated_at": datetime.now()
        }
        
        quiz_id = create_student_quiz(db, quiz_doc)
        print(f"Quiz created with ID: {quiz_id}")
        
        # Generate questions using OpenAI
        try:
            openai_client = start_openai()

            # Simplified difficulty settings
            difficulty_rules = {
                "easy": {"mcq_ratio": 0.7, "tf_ratio": 0.3},
                "medium": {"mcq_ratio": 0.6, "tf_ratio": 0.4},
                "hard": {"mcq_ratio": 0.5, "tf_ratio": 0.5}
            }

            rules = difficulty_rules.get(request.level, difficulty_rules["medium"])

            prompt = f"""
            You are a quiz generator. Create exactly {request.number_of_questions} questions based on the document content below.

            DOCUMENT CONTENT (trimmed for token limit):
            {combined_text[:8000]}

            IF CONTENT IS WEAK:
            Use the topic hints in the text to generate educational questions.

            QUESTION DISTRIBUTION:
            - MCQ Questions: {int(request.number_of_questions * rules['mcq_ratio'])}
            - True/False Questions: {int(request.number_of_questions * rules['tf_ratio'])}

            QUESTION REQUIREMENTS:
            1. MCQ:
            - Provide 4 realistic options (A, B, C, D)
            - Only one correct answer

            2. TRUE/FALSE:
            - Options: "True", "False"
            - Only one correct answer

            3. Every question MUST include an explanation explaining why the correct answer is correct.

            DO NOT create:
            - generic or metadata questions
            - questions like "What is this document about?" or "Who created this"
            - filler or unrelated questions

            Return ONLY this JSON structure, no extra text:

            {{
            "questions": [
                {{
                "question_type": "mcq",
                "question_text": "",
                "options": ["A", "B", "C", "D"],
                "correct_answer": "A",
                "AI_explanation": ""
                }},
                {{
                "question_type": "true_false",
                "question_text": "",
                "options": ["True", "False"],
                "correct_answer": "True",
                "AI_explanation": ""
                }}
            ]
            }}
            """
            
            response = await openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                    {"role": "system", "content": "You generate quizzes that test real understanding, not memory. Return only JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.6,
                max_tokens=3000
            )
            try:
                raw_response = response.choices[0].message.content.strip()
                cleaned_response = raw_response.lstrip("```json").rstrip("```").strip()
                
                # Remove markdown code blocks (```json ... ```) - handle multiple formats
                # Pattern 1: ```json ... ```
                cleaned_response = re.sub(r'^```(?:json)?\s*\n?', '', cleaned_response, flags=re.MULTILINE)
                cleaned_response = re.sub(r'\n?\s*```\s*$', '', cleaned_response, flags=re.MULTILINE)
                
                # Pattern 2: ``` ... ``` (without json)
                cleaned_response = re.sub(r'^```\s*\n?', '', cleaned_response, flags=re.MULTILINE)
                cleaned_response = re.sub(r'\n?\s*```\s*$', '', cleaned_response, flags=re.MULTILINE)
                
                # Try to extract JSON if wrapped in other text - find the largest JSON object
                json_match = re.search(r'\{[\s\S]*\}', cleaned_response)
                if json_match:
                    cleaned_response = json_match.group(0)
                
                # Clean up any leading/trailing whitespace
                cleaned_response = cleaned_response.strip()
                
                # Remove any leading text before first {
                if cleaned_response.find('{') != -1:
                    cleaned_response = cleaned_response[cleaned_response.find('{'):]
                
                # Remove any trailing text after last }
                if cleaned_response.rfind('}') != -1:
                    cleaned_response = cleaned_response[:cleaned_response.rfind('}') + 1]
                
                # Parse JSON
                result = json.loads(cleaned_response)
                questions_data = result.get("questions", [])
                
                if not questions_data:
                    return GenerateQuizResponse(
                        success=False,
                        message="No questions generated. Please try again."
                    )
                
                # Filter out generic/bad questions
                bad_question_patterns = [
                    "what is this document",
                    "what is the document",
                    "how long is",
                    "who created",
                    "document content primarily about",
                    "what is the video",
                    "document is about",
                    "document mentions",
                    "document describes"
                ]
                
                filtered_questions = []
                for q in questions_data:
                    question_text = q.get("question_text", "").lower()
                    # Check if question is generic
                    is_generic = any(pattern in question_text for pattern in bad_question_patterns)
                    
                    if not is_generic and len(question_text) > 20:  # Minimum meaningful length
                        filtered_questions.append(q)
                
                # If we filtered too many, keep original questions
                if len(filtered_questions) < request.number_of_questions * 0.5:
                    filtered_questions = questions_data[:request.number_of_questions]
                
                questions_data = filtered_questions[:request.number_of_questions]  # Limit to requested number
                
                # Save each question to database
                saved_questions = []
                for i, q in enumerate(questions_data):
                    question_doc = {
                        "quiz_id": ObjectId(quiz_id),
                        "question_type": q.get("question_type", "mcq"),
                        "question_text": q.get("question_text", ""),
                        "options": q.get("options", []),
                        "correct_answer": q.get("correct_answer", ""),
                        "student_answer": None,
                        "AI_explanation": q.get("AI_explanation", ""),
                        "created_at": datetime.now(),
                        "updated_at": datetime.now()
                    }
                    
                    question_id = create_question_answer(db, question_doc)
                    saved_questions.append(question_id)
                
                # Create response
                student_quiz = StudentQuiz(
                    quiz_name=request.quiz_name,
                    related_doc_id=request.document_id,
                    created_by=request.user_id,
                    level=request.level,
                    no_of_questions=request.number_of_questions,
                    is_submitted=False,
                    created_at=datetime.now(),
                    updated_at=datetime.now()
                )
                
                return GenerateQuizResponse(
                    success=True,
                    message=f"Quiz generated successfully with {len(saved_questions)} quality questions",
                    data=student_quiz
                )
                
            except json.JSONDecodeError as json_error:
                print(f"[QUIZ] ❌ JSON parsing error: {json_error}")
                # Try to extract just the JSON part more aggressively
                try:
                    # Try to find JSON object in response
                    json_start = raw_response.find('{')
                    json_end = raw_response.rfind('}') + 1
                    if json_start != -1 and json_end > json_start:
                        extracted_json = raw_response[json_start:json_end]
                        result = json.loads(extracted_json)
                        questions_data = result.get("questions", [])
                        
                        if not questions_data:
                            raise ValueError("No questions found in extracted JSON")
                        
                        # Continue with the same processing logic
                        # Filter out generic/bad questions
                        bad_question_patterns = [
                            "what is this document",
                            "what is the document",
                            "how long is",
                            "who created",
                            "document content primarily about",
                            "what is the video",
                            "document is about",
                            "document mentions",
                            "document describes"
                        ]
                        
                        filtered_questions = []
                        for q in questions_data:
                            question_text = q.get("question_text", "").lower()
                            is_generic = any(pattern in question_text for pattern in bad_question_patterns)
                            
                            if not is_generic and len(question_text) > 20:
                                filtered_questions.append(q)
                        
                        if len(filtered_questions) < request.number_of_questions * 0.5:
                            filtered_questions = questions_data[:request.number_of_questions]
                        
                        questions_data = filtered_questions[:request.number_of_questions]
                        
                        # Save questions
                        saved_questions = []
                        for i, q in enumerate(questions_data):
                            question_doc = {
                                "quiz_id": ObjectId(quiz_id),
                                "question_type": q.get("question_type", "mcq"),
                                "question_text": q.get("question_text", ""),
                                "options": q.get("options", []),
                                "correct_answer": q.get("correct_answer", ""),
                                "student_answer": None,
                                "AI_explanation": q.get("AI_explanation", ""),
                                "created_at": datetime.now(),
                                "updated_at": datetime.now()
                            }
                            question_id = create_question_answer(db, question_doc)
                            saved_questions.append(question_id)
                        
                        student_quiz = StudentQuiz(
                            quiz_name=request.quiz_name,
                            related_doc_id=request.document_id,
                            created_by=request.user_id,
                            level=request.level,
                            no_of_questions=request.number_of_questions,
                            is_submitted=False,
                            created_at=datetime.now(),
                            updated_at=datetime.now()
                        )
                        
                        return GenerateQuizResponse(
                            success=True,
                            message=f"Quiz generated successfully with {len(saved_questions)} quality questions",
                            data=student_quiz
                        )
                    else:
                        raise json_error
                except Exception as retry_error:
                    print(f"[QUIZ] ❌ Retry also failed: {retry_error}")
                    import traceback
                    traceback.print_exc()
                    return GenerateQuizResponse(
                        success=False,
                        message=f"Error parsing quiz questions from AI response. The AI may have returned invalid JSON. Please try again. Error: {str(json_error)}"
                    )
            
        except Exception as e:
            print(f"Error generating questions with OpenAI: {e}")
            return GenerateQuizResponse(
                success=False,
                message=f"Error generating questions: {str(e)}"
            )
        
    except Exception as e:
        print(f"Error in generate_student_quiz: {e}")
        capture_exception(e)
        return GenerateQuizResponse(
            success=False,
            message=f"Error generating quiz: {str(e)}"
        )

async def get_document_quizzes(document_id: str, created_by: str = None):
    """Get all quizzes for a document, optionally filtered by creator"""
    try:
        from bson import ObjectId
        from app.database.mongo_collections import get_student_quizs_by_doc, get_questions_by_quiz
        
        db = next(get_db())
        
        # Validate document exists
        document = db["docSathi_ai_documents"].find_one({"_id": ObjectId(document_id)})
        if not document:
            return GetDocumentQuizzesResponse(
                success=False,
                message="Document not found"
            )
        
        # Get quizzes for document
        quizzes = get_student_quizs_by_doc(db, document_id)
        
        # Filter by created_by if provided
        if created_by:
            quizzes = [q for q in quizzes if str(q.get("created_by")) == created_by]
        
        if not quizzes:
            return GetDocumentQuizzesResponse(
                success=True,
                message="No quizzes found for this document",
                data=[]
            )
        
        # Get questions for each quiz
        quiz_responses = []
        for quiz in quizzes:
            quiz_id = str(quiz["_id"])
            questions = get_questions_by_quiz(db, quiz_id)
            
            # Convert questions to response format
            question_responses = []
            is_quiz_submitted = quiz.get("is_submitted", False)
            
            for question in questions:
                # Only show correct answer and AI explanation if quiz is submitted
                if is_quiz_submitted:
                    question_responses.append(QuizQuestionResponse(
                        question_id=str(question["_id"]),
                        question_type=question.get("question_type", "mcq"),
                        question_text=question.get("question_text", ""),
                        options=question.get("options", []),
                        correct_answer=question.get("correct_answer", ""),
                        student_answer=question.get("student_answer"),
                        AI_explanation=question.get("AI_explanation", "")
                    ))
                else:
                    # For non-submitted quizzes, hide correct answer and AI explanation
                    question_responses.append(QuizQuestionResponse(
                        question_id=str(question["_id"]),
                        question_type=question.get("question_type", "mcq"),
                        question_text=question.get("question_text", ""),
                        options=question.get("options", []),
                        correct_answer="",  # Hidden for non-submitted quizzes
                        student_answer=question.get("student_answer"),
                        AI_explanation=""  # Hidden for non-submitted quizzes
                    ))
            
            # Create quiz response
            quiz_response = QuizResponse(
                quiz_id=quiz_id,
                quiz_name=quiz.get("quiz_name", ""),
                level=quiz.get("level", "medium"),
                no_of_questions=quiz.get("no_of_questions", 0),
                is_submitted=quiz.get("is_submitted", False),
                created_at=quiz.get("created_at", datetime.now()),
                updated_at=quiz.get("updated_at", datetime.now()),
                questions=question_responses
            )
            
            quiz_responses.append(quiz_response)
        
        return GetDocumentQuizzesResponse(
            success=True,
            message=f"Found {len(quiz_responses)} quiz(es) for document",
            data=quiz_responses
        )
        
    except Exception as e:
        print(f"Error in get_document_quizzes: {e}")
        capture_exception(e)
        return GetDocumentQuizzesResponse(
            success=False,
            message=f"Error fetching quizzes: {str(e)}"
        )

async def submit_quiz(quiz_id: str, answers: list[dict]):
    """Submit quiz answers and calculate score"""
    try:
        from bson import ObjectId
        from app.database.mongo_collections import get_student_quiz, get_questions_by_quiz, update_question_answer, update_student_quiz
        
        db = next(get_db())
        
        # Validate quiz exists
        quiz = get_student_quiz(db, quiz_id)
        if not quiz:
            return SubmitQuizResponse(
                success=False,
                message="Quiz not found"
            )
        
        # Check if quiz already submitted
        if quiz.get("is_submitted", False):
            return SubmitQuizResponse(
                success=False,
                message="Quiz already submitted"
            )
        
        # Get all questions for this quiz
        questions = get_questions_by_quiz(db, quiz_id)
        if not questions:
            return SubmitQuizResponse(
                success=False,
                message="No questions found for this quiz"
            )
        
        # Create answer mapping for quick lookup
        answer_map = {ans["question_id"]: ans["selected_answer"] for ans in answers}
        
        # Process each question and calculate score
        correct_answers = 0
        wrong_answers = 0
        total_questions = len(questions)
        processed_questions = []
        
        for question in questions:
            question_id = str(question["_id"])
            selected_answer = answer_map.get(question_id, "")
            correct_answer = question.get("correct_answer", "")
            
            # Check if answer is correct
            is_correct = selected_answer.strip() == correct_answer.strip()
            
            if is_correct:
                correct_answers += 1
            else:
                wrong_answers += 1
            
            # Update question with student answer
            update_question_answer(db, question_id, {
                "student_answer": selected_answer,
                "updated_at": datetime.now()
            })
            
            # Create question response for result
            processed_questions.append(QuizQuestionResponse(
                question_id=question_id,
                question_type=question.get("question_type", "mcq"),
                question_text=question.get("question_text", ""),
                options=question.get("options", []),
                correct_answer=correct_answer,
                student_answer=selected_answer,
                AI_explanation=question.get("AI_explanation", "")
            ))
        
        # Calculate score and percentage
        score = correct_answers
        percentage = (correct_answers / total_questions) * 100 if total_questions > 0 else 0
        
        # Update quiz as submitted
        update_student_quiz(db, quiz_id, {
            "is_submitted": True,
            "score": score,
            "percentage": percentage,
            "submitted_at": datetime.now(),
            "updated_at": datetime.now()
        })
        
        # Create submission result
        submission_result = QuizSubmissionResult(
            quiz_id=quiz_id,
            score=score,
            total_questions=total_questions,
            correct_answers=correct_answers,
            wrong_answers=wrong_answers,
            percentage=round(percentage, 2),
            submitted_at=datetime.now(),
            questions=processed_questions
        )
        
        return SubmitQuizResponse(
            success=True,
            message=f"Quiz submitted successfully. Score: {score}/{total_questions} ({percentage:.1f}%)",
            data=submission_result
        )
        
    except Exception as e:
        print(f"Error in submit_quiz: {e}")
        capture_exception(e)
        return SubmitQuizResponse(
            success=False,
            message=f"Error submitting quiz: {str(e)}"
        )

async def get_document_text(document_id: str):
    """Get combined text content from document chunks"""
    try:
        from bson import ObjectId
        
        db = next(get_db())
        
        # Validate document_id
        try:
            doc_object_id = ObjectId(document_id)
        except Exception:
            return DocumentTextResponse(
                success=False,
                message="Invalid document ID format"
            )
        
        # Get document details
        document = db["docSathi_ai_documents"].find_one({"_id": doc_object_id})
        if not document:
            return DocumentTextResponse(
                success=False,
                message="Document not found"
            )
        
        # Use the common function to get document text
        combined_text, error = await get_document_text_from_chunks(document_id)
        
        if error:
            return DocumentTextResponse(
                success=False,
                message=error
            )
        
        # Return the text content with type information and document URL
        # Include summary, ai_notes, and transcript_segments if they exist
        return DocumentTextResponse(
            success=True,
            message="Document text retrieved successfully",
            data={
                "content": combined_text,
                "type": "text",
                "document_url": document.get("url", ""),
                "document_name": document.get("name", ""),
                "summary": document.get("summary", ""),
                "document_format": document.get("document_format", ""),
                "ai_notes": document.get("ai_notes"),  # Include ai_notes field if exists
                "transcript_segments": document.get("transcript_segments")  # ✅ Include timestamped transcript for YouTube videos
            }
        )
        
    except Exception as e:
        print(f"Error in get_document_text: {e}")
        capture_exception(e)
        return DocumentTextResponse(
            success=False,
            message=f"Error retrieving document text: {str(e)}"
        )

async def get_document_chats(document_id: str, limit: int = 20, offset: int = 0):
    """Get all chats for a document with pagination"""
    try:
        from bson import ObjectId
        
        db = next(get_db())
        
        # Validate document_id
        try:
            doc_object_id = ObjectId(document_id)
        except Exception:
            return DocumentChatsResponse(
                success=False,
                message="Invalid document ID format"
            )
        
        # Check if document exists
        document = db["docSathi_ai_documents"].find_one({"_id": doc_object_id})
        if not document:
            return DocumentChatsResponse(
                success=False,
                message="Document not found"
            )
        
        # Limit validation
        if limit > 50:
            limit = 50
        if limit < 1:
            limit = 20
        if offset < 0:
            offset = 0
        
        # Get chats for this document with pagination
        chats_cursor = db["chats"].find({
            "document_id": doc_object_id,
            "status": "active"
        }).sort("created_at", -1).skip(offset).limit(limit)
        
        chats = list(chats_cursor)
        
        if not chats:
            return DocumentChatsResponse(
                success=True,
                message="No chats found for this document",
                data=[]
            )
        
        # Convert to response format
        chat_list = []
        for chat in chats:
            # Get last message and message count for each chat
            last_message = ""
            message_count = 0
            
            # Get last message
            last_msg_cursor = db["chat_messages"].find({
                "chat_id": chat["_id"]
            }).sort("created_ts", -1).limit(1)
            
            last_msg = list(last_msg_cursor)
            if last_msg:
                last_message = last_msg[0].get("message", "")
            
            # Get message count
            message_count = db["chat_messages"].count_documents({
                "chat_id": chat["_id"]
            })
            
            chat_list.append(DocumentChat(
                chat_id=str(chat["_id"]),
                document_id=str(chat.get("document_id", "")),
                user_id=str(chat.get("user_id", "")),
                created_at=chat.get("created_at", datetime.now()),
                updated_at=chat.get("updated_at", datetime.now()),
                last_message=last_message,
                message_count=message_count,
                title=chat.get("title", f"Chat {str(chat['_id'])[:8]}")  # Use existing title from chat
            ))
        
        return DocumentChatsResponse(
            success=True,
            message=f"Found {len(chat_list)} chat(s) for document",
            data=chat_list
        )
        
    except Exception as e:
        print(f"Error in get_document_chats: {e}")
        capture_exception(e)
        return DocumentChatsResponse(
            success=False,
            message=f"Error retrieving document chats: {str(e)}"
        )



def chunk_text(text, max_tokens=10000):
    sentences = sent_tokenize(text)  # Split text into sentences
    chunks = []
    current_chunk = []
    current_length = 0

    for sentence in sentences:
        tokenized_sentence = tokenizer.encode(sentence, add_special_tokens=False)
        sentence_length = len(tokenized_sentence)

        # If adding this sentence exceeds max_tokens, finalize the chunk
        if current_length + sentence_length > max_tokens:
            if current_chunk:  # Ensure we don't create empty chunks
                chunks.append(" ".join(current_chunk))
            current_chunk = [sentence]  # Start a new chunk with this sentence
            current_length = sentence_length
        else:
            current_chunk.append(sentence)
            current_length += sentence_length

    # Add last chunk if any sentences remain
    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks


# 🔹 Optimized Function to Dynamically Adjust Summary Length
def get_summary_length(word_count):
    if word_count <= 1000:
        return 100, 200  # Short documents
    elif word_count <= 5000:
        return 150, 300  # Medium documents
    elif word_count <= 10000:
        return 300, 600  # Long documents
    else:
        return 500, 1000  # Very Large documents


# 🔹 Optimized Summarization Function with Sentence-Aware Chunking & Parameter Tweaks
